"""
CV Template Builder
Generates CV content using modular template approach with focused LLM prompts
"""

import logging
from pathlib import Path
from typing import Dict, List, Optional, Tuple
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

from modules.generation.ollama_client import OllamaClient
from templates.cv_template_structure import get_cv_template, get_section_order


class CVTemplateBuilder:
    """Builds CVs using template approach - LLM generates content, Python handles formatting"""

    def __init__(self, ollama_client: OllamaClient, user_info: dict, prompts_dir: Path):
        self.ollama = ollama_client
        self.user_info = user_info
        self.prompts_dir = prompts_dir
        self.logger = logging.getLogger(__name__)
        self.max_retries = 3

    def generate_cv(self, job, matched_projects: List[Dict]) -> Optional[Document]:
        """
        Generate CV using template approach.

        Args:
            job: Job object with description, title, company
            matched_projects: List of 3 matched projects with metadata

        Returns:
            Document object or None if generation fails
        """
        self.logger.info("Generating CV using template approach...")

        # Step 1: Prepare static content (no LLM needed)
        static_content = self._prepare_static_content(matched_projects)

        # Step 2: Generate dynamic content sections with focused prompts
        dynamic_content = self._generate_dynamic_content(job, matched_projects)

        if not dynamic_content:
            self.logger.error("Failed to generate dynamic content")
            return None

        # Step 3: Merge all content
        all_content = {**static_content, **dynamic_content}

        # Step 4: Fill template and create Word document
        cv_doc = self._create_word_document(all_content)

        return cv_doc

    def _prepare_static_content(self, matched_projects: List[Dict]) -> Dict[str, str]:
        """Prepare content that doesn't need LLM generation"""
        content = {}

        # User info
        content['NAME'] = self.user_info.get('name', '')
        content['EMAIL'] = self.user_info.get('email', '')
        content['PHONE'] = self.user_info.get('phone', '')
        content['LINKEDIN'] = self.user_info.get('linkedin', '')
        content['LOCATION'] = self.user_info.get('location', '')

        # Skills - Templated (no LLM generation)
        # Fixed categories to prevent LLM circular logic like "Machine Learning & AI: Machine Learning"
        content['SKILLS_CATEGORIZED'] = """Programming & Tools: Python, SQL, Git, Jupyter
ML/AI Frameworks: TensorFlow, Keras, PyTorch, Scikit-learn, XGBoost
Deep Learning: CNN, LSTM, Transformers (BERT), Computer Vision (YOLOv8, OpenCV)
Data Engineering: API Integration, Data Pipelines, FFmpeg, Pandas, NumPy
Specializations: NLP, Sentiment Analysis, Automated Systems, Model Optimization"""

        # Professional Experience - Templated (no LLM generation)
        content['PROFESSIONAL_EXPERIENCE'] = """Data Analytics Intern
RCCG Department of Public Health | January 2024 – November 2024

• Processed 500+ monthly financial transactions maintaining 99.8% accuracy
• Analyzed payment patterns to identify trends for health outreach programs
• Generated weekly reports supporting data-driven clinic operations
• Collaborated with medical staff to optimize workflows, reducing processing time by 25%"""

        # Education - Include BOTH degrees with proper formatting (templated, not LLM-generated)
        education_lines = []

        # MSc (current degree)
        education_lines.append("Master of Science in Data Science (MSc)")
        education_lines.append("University of Hertfordshire, UK | Expected Graduation: June 2026")
        education_lines.append("Relevant Coursework: Machine Learning, Deep Learning, NLP, Big Data Analytics, Time Series Analysis")
        education_lines.append("")  # Blank line

        # BSc (previous degree)
        education_lines.append("Bachelor of Science in Anatomy (BSc)")
        education_lines.append("Bowen University | Graduated: May 2019")
        education_lines.append("Relevant Coursework: Biostatistics, Research Methodology, Data Analysis, Quantitative Methods")

        content['EDUCATION'] = "\n".join(education_lines)

        # Certifications (optional)
        certs = self.user_info.get('certifications', [])
        if certs:
            content['CERTIFICATIONS'] = "CERTIFICATIONS\n" + "\n".join(f"• {cert}" for cert in certs)
        else:
            content['CERTIFICATIONS'] = ""

        # Project titles and technologies (from matched projects)
        for i, project in enumerate(matched_projects[:3], 1):
            content[f'PROJECT_{i}_TITLE'] = project['title']
            content[f'PROJECT_{i}_TECH'] = ", ".join(project['technologies'])

        return content

    def _generate_dynamic_content(self, job, matched_projects: List[Dict]) -> Optional[Dict[str, str]]:
        """Generate content sections using focused LLM prompts"""
        content = {}

        # Section 1: Professional Summary
        summary = self._generate_section(
            section_name='professional_summary',
            prompt_file='professional_summary.txt',
            job=job,
            matched_projects=matched_projects
        )
        if not summary:
            return None
        content['PROFESSIONAL_SUMMARY'] = summary

        # Skills are now templated in _prepare_static_content() - no LLM generation needed
        # This prevents circular logic like "Machine Learning & AI: Machine Learning"

        # Section 2-4: Project bullets (one section per project)
        for i, project in enumerate(matched_projects[:3], 1):
            bullets = self._generate_project_bullets(job, project, i)
            if not bullets:
                return None
            content[f'PROJECT_{i}_BULLETS'] = bullets

        return content

    def _generate_section(self, section_name: str, prompt_file: str, job, matched_projects: List[Dict]) -> Optional[str]:
        """Generate a single section with retry logic"""

        for attempt in range(1, self.max_retries + 1):
            try:
                self.logger.info(f"Generating {section_name} (attempt {attempt}/{self.max_retries})")

                # Load prompt template
                prompt_path = self.prompts_dir / 'cv_sections' / prompt_file
                with open(prompt_path, 'r', encoding='utf-8') as f:
                    prompt_template = f.read()

                # Prepare context
                context = self._prepare_context(job, matched_projects)

                # Fill prompt
                prompt = prompt_template.format(**context)

                # Generate with focused parameters
                temperature = 0.3 if attempt == 1 else 0.2  # Lower temp on retry
                max_tokens = 300 if section_name == 'professional_summary' else 200

                content = self.ollama.generate_text(
                    prompt=prompt,
                    temperature=temperature,
                    max_tokens=max_tokens
                )

                if not content:
                    self.logger.warning(f"Empty content for {section_name}")
                    continue

                # Validate section
                is_valid, issues = self._validate_section(content, section_name)

                if is_valid:
                    self.logger.info(f"✓ {section_name} generated successfully")
                    return content.strip()
                else:
                    self.logger.warning(f"✗ {section_name} validation failed: {', '.join(issues)}")

            except Exception as e:
                self.logger.error(f"Error generating {section_name}: {e}")

        self.logger.error(f"Failed to generate {section_name} after {self.max_retries} attempts")
        return None

    def _generate_project_bullets(self, job, project: Dict, project_num: int) -> Optional[str]:
        """Generate bullet points for a single project"""

        for attempt in range(1, self.max_retries + 1):
            try:
                self.logger.info(f"Generating project {project_num} bullets (attempt {attempt}/{self.max_retries})")

                # Load prompt
                prompt_path = self.prompts_dir / 'cv_sections' / 'project_bullets.txt'
                with open(prompt_path, 'r', encoding='utf-8') as f:
                    prompt_template = f.read()

                # Prepare context
                context = {
                    'job_description': job.description[:1000],
                    'project_title': project['title'],
                    'project_technologies': ", ".join(project['technologies']),
                    'project_description': project['description'],
                    'relevance_context': project.get('relevance_context', '')
                }

                # Fill prompt
                prompt = prompt_template.format(**context)

                # Generate
                temperature = 0.4 if attempt == 1 else 0.3
                bullets = self.ollama.generate_text(
                    prompt=prompt,
                    temperature=temperature,
                    max_tokens=250
                )

                if not bullets:
                    continue

                # Validate bullets
                is_valid, issues = self._validate_bullets(bullets)

                if is_valid:
                    self.logger.info(f"✓ Project {project_num} bullets generated successfully")
                    return bullets.strip()
                else:
                    self.logger.warning(f"✗ Project {project_num} bullets validation failed: {', '.join(issues)}")

            except Exception as e:
                self.logger.error(f"Error generating project {project_num} bullets: {e}")

        self.logger.error(f"Failed to generate project {project_num} bullets after {self.max_retries} attempts")
        return None

    def _prepare_context(self, job, matched_projects: List[Dict]) -> Dict[str, str]:
        """Prepare context for prompt templates"""

        # User info summary
        user_info_text = f"""
Name: {self.user_info.get('name', '')}
Education: {self.user_info.get('current_education', '')}
Skills: {', '.join(self.user_info.get('skills', [])[:15])}
Experience: {self.user_info.get('experience_summary', 'Recent graduate')}
        """.strip()

        # Projects summary
        projects_summary = "\n".join([
            f"- {p['title']} (Relevance: {p.get('score', 0)}/10): {p['description'][:150]}"
            for p in matched_projects[:3]
        ])

        # All skills
        all_skills = ', '.join(self.user_info.get('skills', []))

        return {
            'job_description': job.description[:1500],
            'user_info': user_info_text,
            'projects_summary': projects_summary,
            'all_skills': all_skills
        }

    def _validate_section(self, content: str, section_name: str) -> Tuple[bool, List[str]]:
        """Validate generated section content"""
        issues = []

        # Common validation for all sections
        content_lower = content.lower()

        # Check for meta-commentary
        meta_phrases = ['here is', 'here are', 'i have generated', 'i have created']
        if any(phrase in content_lower for phrase in meta_phrases):
            issues.append("Contains meta-commentary")

        # Check for placeholders
        placeholder_patterns = [r'\[.*?\]', r'\{\{.*?\}\}', 'TBD', 'N/A', 'not provided']
        import re
        for pattern in placeholder_patterns:
            if re.search(pattern, content, re.IGNORECASE):
                issues.append(f"Contains placeholder: {pattern}")

        # Check for AI clichés
        cliches = ['leveraged', 'spearheaded', 'passion for', 'team player', 'cutting-edge', 'fast-paced']
        found_cliches = [c for c in cliches if c in content_lower]
        if found_cliches:
            issues.append(f"Contains clichés: {', '.join(found_cliches)}")

        # Section-specific validation
        if section_name == 'professional_summary':
            sentences = content.split('.')
            if len(sentences) < 2 or len(sentences) > 4:
                issues.append("Summary should be 2-3 sentences")

        elif section_name == 'skills_categorized':
            if ':' not in content:
                issues.append("Skills must be categorized with format 'Category: skills'")
            lines = [l for l in content.split('\n') if l.strip()]
            if len(lines) < 2 or len(lines) > 4:
                issues.append("Should have 2-4 skill categories")

        return len(issues) == 0, issues

    def _validate_bullets(self, bullets: str) -> Tuple[bool, List[str]]:
        """Validate project bullet points"""
        issues = []

        # Check bullet count
        bullet_lines = [line for line in bullets.split('\n') if line.strip().startswith('•')]
        if len(bullet_lines) < 2:
            issues.append("Must have at least 2 bullet points")
        elif len(bullet_lines) > 3:
            issues.append("Must have at most 3 bullet points")

        # Check for placeholders
        import re
        placeholder_patterns = [r'\[.*?\]', r'\{\{.*?\}\}']
        for pattern in placeholder_patterns:
            if re.search(pattern, bullets):
                issues.append(f"Contains placeholder: {pattern}")

        # Check for clichés
        bullets_lower = bullets.lower()
        cliches = ['leveraged', 'spearheaded', 'passion for', 'cutting-edge']
        found_cliches = [c for c in cliches if c in bullets_lower]
        if found_cliches:
            issues.append(f"Contains clichés: {', '.join(found_cliches)}")

        # Check for meta-commentary
        if any(phrase in bullets_lower for phrase in ['here is', 'here are', 'bullet point']):
            issues.append("Contains meta-commentary")

        return len(issues) == 0, issues

    def _create_word_document(self, content: Dict[str, str]) -> Document:
        """Create formatted Word document from content"""
        doc = Document()

        # Set margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(0.5)
            section.bottom_margin = Inches(0.5)
            section.left_margin = Inches(0.7)
            section.right_margin = Inches(0.7)

        # Name (16pt, bold, centered)
        name_para = doc.add_paragraph(content.get('NAME', 'Unknown'))
        name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if name_para.runs:
            name_run = name_para.runs[0]
            name_run.font.size = Pt(16)
            name_run.bold = True

        # Contact info - Line 1: Location | Phone | Email (10pt, centered)
        contact_line1 = f"{content.get('LOCATION', '')} | {content.get('PHONE', '')} | {content.get('EMAIL', '')}"
        contact_para1 = doc.add_paragraph(contact_line1)
        contact_para1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if contact_para1.runs:
            contact_para1.runs[0].font.size = Pt(10)

        # Contact info - Line 2: LinkedIn | GitHub (10pt, centered)
        github = self.user_info.get('github', '')
        contact_line2 = f"{content.get('LINKEDIN', '')}"
        if github:
            contact_line2 += f" | {github}"
        contact_para2 = doc.add_paragraph(contact_line2)
        contact_para2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if contact_para2.runs:
            contact_para2.runs[0].font.size = Pt(10)

        # Professional Summary
        self._add_section(doc, "PROFESSIONAL SUMMARY", content['PROFESSIONAL_SUMMARY'])

        # Technical Skills
        self._add_section(doc, "TECHNICAL SKILLS", content['SKILLS_CATEGORIZED'])

        # Key Projects
        header_para = doc.add_paragraph("KEY PROJECTS")
        if header_para.runs:
            header_para.runs[0].font.size = Pt(12)
            header_para.runs[0].bold = True

        # Only iterate over projects that actually exist in content
        num_projects = sum(1 for i in range(1, 4) if f'PROJECT_{i}_TITLE' in content and content[f'PROJECT_{i}_TITLE'])

        for i in range(1, num_projects + 1):
            # Project title
            title = content.get(f'PROJECT_{i}_TITLE', '')
            if not title:
                continue

            title_para = doc.add_paragraph(title)
            if title_para.runs:
                title_para.runs[0].font.size = Pt(11)
                title_para.runs[0].bold = True

            # Technologies
            tech = content.get(f'PROJECT_{i}_TECH', '')
            tech_para = doc.add_paragraph(f"Technologies: {tech}")
            if tech_para.runs:
                tech_para.runs[0].font.size = Pt(11)

            # Bullets
            bullets = content.get(f'PROJECT_{i}_BULLETS', '')
            for bullet_line in bullets.split('\n'):
                if bullet_line.strip():
                    bullet_para = doc.add_paragraph(bullet_line.strip())
                    if bullet_para.runs:
                        bullet_para.runs[0].font.size = Pt(11)

            # Spacing
            doc.add_paragraph()

        # Professional Experience (before Education)
        self._add_section(doc, "PROFESSIONAL EXPERIENCE", content['PROFESSIONAL_EXPERIENCE'])

        # Education
        self._add_section(doc, "EDUCATION", content['EDUCATION'])

        # Certifications (if any)
        if content.get('CERTIFICATIONS'):
            cert_lines = content['CERTIFICATIONS'].split('\n')
            if cert_lines and cert_lines[0]:
                cert_header = doc.add_paragraph(cert_lines[0])
                if cert_header.runs:
                    cert_header.runs[0].font.size = Pt(12)
                    cert_header.runs[0].bold = True
            for cert_line in cert_lines[1:]:
                if cert_line.strip():
                    cert_para = doc.add_paragraph(cert_line.strip())
                    if cert_para.runs:
                        cert_para.runs[0].font.size = Pt(11)

        return doc

    def _add_section(self, doc: Document, header: str, content: str):
        """Add a section with header and content"""
        # Header (12pt, bold)
        header_para = doc.add_paragraph(header)
        if header_para.runs:
            header_para.runs[0].font.size = Pt(12)
            header_para.runs[0].bold = True

        # Content (11pt)
        for line in content.split('\n'):
            if line.strip():
                para = doc.add_paragraph(line.strip())
                if para.runs:
                    para.runs[0].font.size = Pt(11)

        # Spacing
        doc.add_paragraph()

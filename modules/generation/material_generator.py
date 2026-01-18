"""
Material Generator

Generates ATS-optimized CVs and cover letters for job applications.
"""

import os
from pathlib import Path
from typing import Dict, Any, Optional, Tuple
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

from modules.generation.ollama_client import get_ollama_client
from modules.generation.project_matcher import get_project_matcher
from modules.scraping.job_models import JobPosting
from modules.core.logger import get_logger
from modules.utils.helpers import sanitize_filename
from modules.utils.cv_validator import CVValidator
from modules.utils.cv_fixer import get_cv_fixer


class MaterialGenerator:
    """Generates CVs and cover letters"""

    def __init__(
        self,
        user_info: Dict[str, Any],
        output_dir: str = "workspace/generated_materials",
    ):
        """
        Initialize material generator

        Args:
            user_info: User information dictionary
            output_dir: Output directory for generated materials
        """
        self.user_info = user_info
        self.output_dir = Path(output_dir)
        self.logger = get_logger()
        self.ollama = get_ollama_client()
        self.project_matcher = get_project_matcher()
        self.validator = CVValidator()
        self.fixer = get_cv_fixer()

        # Create output directories
        (self.output_dir / "cvs").mkdir(parents=True, exist_ok=True)
        (self.output_dir / "cover_letters").mkdir(parents=True, exist_ok=True)
        (self.output_dir / "failed_cvs").mkdir(parents=True, exist_ok=True)

    def format_user_info(self) -> str:
        """Format user information with proper defaults"""

        # Handle missing graduation date
        grad_date = self.user_info.get('graduation_date')
        if not grad_date:
            # Infer from current education if possible
            current_edu = self.user_info.get('current_education', '')
            if 'expected' in current_edu.lower() or 'msc' in current_edu.lower():
                grad_date = 'Expected June 2026'  # Default for current students
            else:
                grad_date = '[OMIT - not provided]'

        # Handle missing previous education
        prev_edu = self.user_info.get('previous_education', '')
        if not prev_edu or len(prev_edu.strip()) == 0:
            prev_edu = '[OMIT - not provided]'

        info = f"""
Name: {self.user_info.get('name', 'Unknown')}
Email: {self.user_info.get('email', '')}
Phone: {self.user_info.get('phone', '')}
Location: {self.user_info.get('location', '')}
LinkedIn: {self.user_info.get('linkedin', '')}
GitHub: {self.user_info.get('github', '')}

Education:
Current: {self.user_info.get('current_education', '')}
Graduation Date: {grad_date}
Previous: {prev_edu}

Experience: {self.user_info.get('years_experience', {}).get('total', 0)} years total

Skills: {', '.join(self.user_info.get('skills', []))}

IMPORTANT: Any field marked [OMIT - not provided] should be completely excluded from the CV output.
"""
        return info

    def _build_correction_prompt(self, validation_result: Dict[str, Any]) -> str:
        """Build specific correction instructions based on validation issues"""
        corrections = []
        corrections.append("CRITICAL CORRECTIONS NEEDED:")
        corrections.append("")

        for issue in validation_result["critical_issues"]:
            if "meta-commentary" in issue.lower():
                corrections.append("• START IMMEDIATELY with the candidate's name. NO introductory text like 'Here is...'")
            elif "relevance score" in issue.lower():
                corrections.append("• REMOVE all relevance scores (X/10) from the output. They should not appear in the CV.")
            elif "placeholder" in issue.lower():
                corrections.append("• REMOVE all placeholder text [like this], TBD, N/A. Use only complete information.")
            elif "missing" in issue.lower() and "section" in issue.lower():
                corrections.append(f"• INCLUDE the missing section: {issue}")
            elif "contact" in issue.lower():
                corrections.append("• ENSURE contact information (email, phone) is included after the name")
            elif "date" in issue.lower():
                corrections.append(f"• FIX date mismatch: {issue}. Use ONLY dates from provided user info.")
            elif "project" in issue.lower():
                corrections.append("• INCLUDE ALL 3 projects. Each needs 2-3 bullet points.")
            else:
                corrections.append(f"• {issue}")

        corrections.append("")
        corrections.append("Regenerate the CV fixing these specific issues.")

        return "\n".join(corrections)

    def generate_cv(
        self, job: JobPosting, matched_projects: Dict[str, Any]
    ) -> Optional[Tuple[str, str]]:
        """
        Generate ATS-optimized CV

        Args:
            job: Job posting
            matched_projects: Matched projects from project matcher

        Returns:
            Tuple of (cv_text, cv_file_path) or None if failed
        """
        self.logger.info(f"Generating CV for {job.company} - {job.title}")

        # Read prompt template
        prompt_path = Path("prompts/cv_generation_ats.txt")
        if not prompt_path.exists():
            self.logger.error("CV generation prompt not found")
            return None

        with open(prompt_path, "r", encoding="utf-8") as f:
            prompt_template = f.read()

        # Format matched projects - ordered by relevance score
        project_details = []
        project_scores = matched_projects.get("project_scores", [])[:3]

        for i, score_data in enumerate(project_scores, 1):
            project_id = score_data.get("project_id")
            project_name = score_data.get("project_name", "Unknown")
            relevance_score = score_data.get("score", 0)
            reasoning = score_data.get("reasoning", "")

            details = self.project_matcher.get_project_details(project_id)
            if details:
                # Add relevance ranking to help LLM prioritize
                ranked_detail = f"PROJECT #{i} (Relevance Score: {relevance_score}/10)\n"
                ranked_detail += f"Why relevant: {reasoning}\n\n"
                ranked_detail += details
                project_details.append(ranked_detail)

        matched_projects_text = "\n\n" + "="*60 + "\n\n".join(project_details)

        # Prepare prompt variables
        variables = {
            "company": job.company,
            "title": job.title,
            "location": job.location,
            "job_description": job.description[:2000],  # Limit length
            "user_info": self.format_user_info(),
            "matched_projects": matched_projects_text,
        }

        # Substitute variables
        prompt = prompt_template
        for key, value in variables.items():
            prompt = prompt.replace(f"{{{{{key}}}}}", str(value))

        # Generate CV with validation and retries
        self.logger.info("Generating CV content with Llama...")

        cv_text = None
        validation_result = None
        max_retries = 3

        for attempt in range(max_retries):
            if attempt > 0:
                self.logger.info(f"Retry attempt {attempt + 1}/{max_retries}")

            # Attempt 0: Initial generation
            if attempt == 0:
                cv_text = self.ollama.generate_text(
                    prompt,
                    temperature=0.5,
                    max_tokens=2000,
                )

            # Attempt 1: Try targeted fixes first (faster, more reliable)
            elif attempt == 1 and cv_text and validation_result:
                can_fix, fixable_issues = self.fixer.can_fix_issues(
                    validation_result["critical_issues"]
                )

                if can_fix:
                    self.logger.info("Attempting targeted section fixes...")
                    fixed_cv = self.fixer.fix_cv(
                        cv_text,
                        validation_result["critical_issues"],
                        self.user_info,
                        {
                            "company": job.company,
                            "title": job.title,
                            "description": job.description,
                        },
                        matched_projects,
                    )

                    if fixed_cv:
                        cv_text = fixed_cv
                    else:
                        # Fall back to full regeneration
                        self.logger.info("Targeted fixes failed, doing full regeneration...")
                        correction_instructions = self._build_correction_prompt(validation_result)
                        prompt = prompt + "\n\n" + correction_instructions
                        cv_text = self.ollama.generate_text(
                            prompt,
                            temperature=0.4,
                            max_tokens=2000,
                        )
                else:
                    # Issues not fixable with sections, do full regeneration
                    correction_instructions = self._build_correction_prompt(validation_result)
                    prompt = prompt + "\n\n" + correction_instructions
                    cv_text = self.ollama.generate_text(
                        prompt,
                        temperature=0.4,
                        max_tokens=2000,
                    )

            # Attempt 2+: Full regeneration with stricter params
            else:
                correction_instructions = self._build_correction_prompt(validation_result)
                prompt = prompt + "\n\n" + correction_instructions
                cv_text = self.ollama.generate_text(
                    prompt,
                    temperature=0.3,
                    max_tokens=2000,
                )

            if not cv_text:
                self.logger.error(f"Failed to generate CV on attempt {attempt + 1}")
                continue

            # Validate the generated/fixed CV
            validation_result = self.validator.validate(cv_text, self.user_info)

            if validation_result["valid"]:
                self.logger.info(f"✓ CV passed validation (AI score: {validation_result['ai_score']}/100)")
                if validation_result["warnings"]:
                    self.logger.warning(f"CV has {len(validation_result['warnings'])} warnings")
                break
            else:
                self.logger.warning(f"❌ CV validation failed on attempt {attempt + 1}")
                self.logger.warning(f"Issues: {', '.join(validation_result['critical_issues'][:3])}")

        # If all retries failed, save for manual review
        if not validation_result or not validation_result["valid"]:
            self.logger.error("CV generation failed after all retries")
            if cv_text:
                # Save failed CV for debugging
                safe_company = sanitize_filename(job.company)
                failed_path = self.output_dir / "failed_cvs" / f"FAILED_{safe_company}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.txt"
                with open(failed_path, "w", encoding="utf-8") as f:
                    f.write("=== FAILED CV GENERATION ===\n\n")
                    f.write(cv_text)
                    f.write("\n\n=== VALIDATION REPORT ===\n\n")
                    f.write(self.validator.format_validation_report(validation_result))
                self.logger.error(f"Failed CV saved to: {failed_path}")
            return None

        # Save as .docx
        try:
            safe_company = sanitize_filename(job.company)
            filename = f"{self.user_info['name'].replace(' ', '_')}_CV_{safe_company}_ATS"
            cv_path = self.output_dir / "cvs" / f"{filename}.docx"

            # Create professional Word document
            doc = Document()

            # Set margins (standard professional CV margins)
            sections = doc.sections
            for section in sections:
                section.top_margin = Inches(0.5)
                section.bottom_margin = Inches(0.5)
                section.left_margin = Inches(0.7)
                section.right_margin = Inches(0.7)

            # Process CV content with professional formatting
            lines = cv_text.split("\n")
            for line in lines:
                line = line.strip()

                if not line:
                    # Empty line - add small spacing
                    doc.add_paragraph()
                    continue

                # Detect section headers (all caps or starts with common headers)
                is_header = (
                    line.isupper() or
                    line.startswith("EDUCATION") or
                    line.startswith("PROJECTS") or
                    line.startswith("TECHNICAL SKILLS") or
                    line.startswith("EXPERIENCE") or
                    line.startswith("SKILLS") or
                    "────" in line  # Separator lines
                )

                # Detect name (first non-empty line, typically)
                is_name = (
                    len(doc.paragraphs) == 0 and
                    not line.startswith("•") and
                    not ":" in line[:20]
                )

                # Skip separator lines
                if "────" in line or "===" in line:
                    doc.add_paragraph()
                    continue

                # Add paragraph with appropriate formatting
                if is_name:
                    # Name - large, bold
                    p = doc.add_paragraph()
                    run = p.add_run(line)
                    run.font.name = "Calibri"
                    run.font.size = Pt(16)
                    run.font.bold = True
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

                elif is_header:
                    # Section header - bold, slightly larger
                    doc.add_paragraph()  # Add space before header
                    p = doc.add_paragraph()
                    run = p.add_run(line)
                    run.font.name = "Calibri"
                    run.font.size = Pt(12)
                    run.font.bold = True
                    p.space_after = Pt(6)

                elif line.startswith("•"):
                    # Bullet point
                    p = doc.add_paragraph(line[1:].strip(), style='List Bullet')
                    for run in p.runs:
                        run.font.name = "Calibri"
                        run.font.size = Pt(11)
                    p.paragraph_format.left_indent = Inches(0.25)
                    p.space_after = Pt(3)

                elif "|" in line and len(line) < 100:
                    # Contact info or inline details - centered
                    p = doc.add_paragraph()
                    run = p.add_run(line)
                    run.font.name = "Calibri"
                    run.font.size = Pt(10)
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

                else:
                    # Regular text
                    p = doc.add_paragraph(line)
                    for run in p.runs:
                        run.font.name = "Calibri"
                        run.font.size = Pt(11)
                    p.space_after = Pt(3)

            # Save document
            doc.save(cv_path)

            self.logger.info(f"✓ CV saved: {cv_path}")

            return (cv_text, str(cv_path), validation_result)

        except Exception as e:
            self.logger.error(f"Error saving CV: {e}")
            return None

    def generate_cover_letter(
        self, job: JobPosting, matched_projects: Dict[str, Any]
    ) -> Optional[Tuple[str, str]]:
        """
        Generate cover letter

        Args:
            job: Job posting
            matched_projects: Matched projects from project matcher

        Returns:
            Tuple of (cover_letter_text, file_path) or None if failed
        """
        self.logger.info(f"Generating cover letter for {job.company} - {job.title}")

        # Read prompt template
        prompt_path = Path("prompts/cover_letter_generation.txt")
        if not prompt_path.exists():
            self.logger.error("Cover letter generation prompt not found")
            return None

        with open(prompt_path, "r", encoding="utf-8") as f:
            prompt_template = f.read()

        # Get top project
        top_project_id = matched_projects.get("top_projects", [])[0] if matched_projects.get("top_projects") else None
        top_project_text = ""
        if top_project_id:
            top_project_text = self.project_matcher.get_project_details(top_project_id)

        # Prepare prompt variables
        variables = {
            "company": job.company,
            "title": job.title,
            "location": job.location,
            "job_description": job.description[:1500],
            "user_info": self.format_user_info(),
            "top_project": top_project_text,
        }

        # Substitute variables
        prompt = prompt_template
        for key, value in variables.items():
            prompt = prompt.replace(f"{{{{{key}}}}}", str(value))

        # Generate cover letter
        self.logger.info("Generating cover letter content with Llama...")

        cover_letter_text = self.ollama.generate_text(
            prompt,
            temperature=0.6,  # Slightly more creative
            max_tokens=800,
        )

        if not cover_letter_text:
            self.logger.error("Failed to generate cover letter")
            return None

        # Save as .docx
        try:
            safe_company = sanitize_filename(job.company)
            filename = f"{safe_company}_Cover_Letter"
            cl_path = self.output_dir / "cover_letters" / f"{filename}.docx"

            # Create Word document
            doc = Document()

            # Set margins
            sections = doc.sections
            for section in sections:
                section.top_margin = Inches(1)
                section.bottom_margin = Inches(1)
                section.left_margin = Inches(1)
                section.right_margin = Inches(1)

            # Add cover letter content
            for line in cover_letter_text.split("\n"):
                paragraph = doc.add_paragraph(line)
                # Use standard font
                for run in paragraph.runs:
                    run.font.name = "Calibri"
                    run.font.size = Pt(11)

            # Save document
            doc.save(cl_path)

            self.logger.info(f"✓ Cover letter saved: {cl_path}")

            return (cover_letter_text, str(cl_path))

        except Exception as e:
            self.logger.error(f"Error saving cover letter: {e}")
            return None

    def generate_materials(
        self, job: JobPosting
    ) -> Optional[Dict[str, Any]]:
        """
        Generate all materials for a job application

        Args:
            job: Job posting

        Returns:
            Dictionary containing generated materials and metadata
        """
        self.logger.info("=" * 60)
        self.logger.info(f"Generating materials for: {job.company} - {job.title}")
        self.logger.info("=" * 60)

        # Step 1: Match projects
        matched_projects = self.project_matcher.match_projects(job)

        # Step 2: Generate CV
        cv_result = self.generate_cv(job, matched_projects)
        if not cv_result:
            self.logger.error("CV generation failed")
            return None

        cv_text, cv_path, cv_validation = cv_result

        # Step 3: Generate cover letter
        cl_result = self.generate_cover_letter(job, matched_projects)
        if not cl_result:
            self.logger.error("Cover letter generation failed")
            return None

        cl_text, cl_path = cl_result

        # Return all materials
        return {
            "job_id": job.job_id,
            "company": job.company,
            "title": job.title,
            "matched_projects": matched_projects,
            "cv": {
                "text": cv_text,
                "path": cv_path,
                "validation": cv_validation,
            },
            "cover_letter": {
                "text": cl_text,
                "path": cl_path,
            },
            "generated_at": datetime.now().isoformat(),
        }


def create_material_generator(user_info: Dict[str, Any]) -> MaterialGenerator:
    """
    Create material generator

    Args:
        user_info: User information dictionary

    Returns:
        MaterialGenerator instance
    """
    return MaterialGenerator(user_info)

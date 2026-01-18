"""
Cover Letter Template Builder
Generates cover letters using modular template approach with company research
"""

import logging
from pathlib import Path
from typing import Dict, List, Optional, Tuple
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches

from modules.generation.ollama_client import OllamaClient


class CoverLetterTemplateBuilder:
    """Builds cover letters using template approach with web research"""

    def __init__(self, ollama_client: OllamaClient, user_info: dict, prompts_dir: Path):
        self.ollama = ollama_client
        self.user_info = user_info
        self.prompts_dir = prompts_dir
        self.logger = logging.getLogger(__name__)
        self.max_retries = 3

    def generate_cover_letter(self, job, company_research: Dict[str, str], matched_projects: List[Dict]) -> Optional[Document]:
        """
        Generate cover letter using template approach with company research.

        Args:
            job: Job object with description, title, company
            company_research: Dict with company info from web search
            matched_projects: List of matched projects

        Returns:
            Document object or None if generation fails
        """
        self.logger.info("Generating cover letter using template approach...")

        # Prepare static content (header)
        static_content = self._prepare_header()

        # Generate dynamic sections
        opening = self._generate_opening(job, company_research)
        if not opening:
            self.logger.warning("Opening generation failed, using fallback")
            opening = self._fallback_opening(job)

        achievements = self._generate_achievements(job, company_research, matched_projects)
        if not achievements:
            self.logger.warning("Achievements generation failed, using fallback")
            achievements = self._fallback_achievements(matched_projects)

        closing = self._generate_closing(job, company_research)
        if not closing:
            self.logger.warning("Closing generation failed, using fallback")
            closing = self._fallback_closing(job)

        # Create Word document
        doc = self._create_word_document(static_content, opening, achievements, closing, job)

        return doc

    def _prepare_header(self) -> Dict[str, str]:
        """Prepare static header content"""
        github = self.user_info.get('github', '')
        return {
            'NAME': self.user_info.get('name', ''),
            'LOCATION': self.user_info.get('location', ''),
            'PHONE': self.user_info.get('phone', ''),
            'EMAIL': self.user_info.get('email', ''),
            'LINKEDIN': self.user_info.get('linkedin', ''),
            'GITHUB': github,
            'DATE': datetime.now().strftime('%B %d, %Y')
        }

    def _generate_opening(self, job, company_research: Dict[str, str]) -> Optional[str]:
        """Generate personalized opening paragraph based on company research"""

        for attempt in range(1, self.max_retries + 1):
            try:
                self.logger.info(f"Generating opening paragraph (attempt {attempt}/{self.max_retries})")

                # Load prompt
                prompt_path = self.prompts_dir / 'cover_letter_sections' / 'opening.txt'
                with open(prompt_path, 'r', encoding='utf-8') as f:
                    prompt_template = f.read()

                # Prepare context
                context = {
                    'job_title': job.title,
                    'company_name': job.company,
                    'company_mission': company_research.get('mission', 'N/A'),
                    'company_products': company_research.get('products', 'N/A'),
                    'company_news': company_research.get('recent_news', 'N/A'),
                    'tech_stack': company_research.get('tech_stack', 'N/A'),
                    'job_description': job.description[:800]
                }

                prompt = prompt_template.format(**context)

                # Generate with low temperature for consistency
                opening = self.ollama.generate_text(
                    prompt=prompt,
                    temperature=0.4,
                    max_tokens=200
                )

                if opening and len(opening.strip()) > 50:
                    self.logger.info("✓ Opening paragraph generated successfully")
                    return opening.strip()

            except Exception as e:
                self.logger.error(f"Error generating opening: {e}")

        return None

    def _generate_achievements(self, job, company_research: Dict[str, str], matched_projects: List[Dict]) -> Optional[str]:
        """Generate achievement paragraph based on job requirements and company research"""

        for attempt in range(1, self.max_retries + 1):
            try:
                self.logger.info(f"Generating achievements (attempt {attempt}/{self.max_retries})")

                # Load prompt
                prompt_path = self.prompts_dir / 'cover_letter_sections' / 'achievements.txt'
                with open(prompt_path, 'r', encoding='utf-8') as f:
                    prompt_template = f.read()

                # Prepare work experience
                work_experience = """Data Analytics Intern at RCCG Department of Public Health (January 2024 – November 2024):
- Processed 500+ monthly financial transactions maintaining 99.8% accuracy
- Analyzed payment patterns to identify trends for health outreach programs
- Generated weekly reports supporting data-driven clinic operations
- Collaborated with medical staff to optimize workflows, reducing processing time by 25%"""

                # Prepare projects
                projects_text = "\n\n".join([
                    f"{p['title']}\n{p['description']}\nTechnologies: {', '.join(p['technologies'])}"
                    for p in matched_projects[:3]
                ])

                context = {
                    'job_title': job.title,
                    'job_description': job.description[:1000],
                    'company_name': job.company,
                    'tech_stack': company_research.get('tech_stack', 'N/A'),
                    'work_experience': work_experience,
                    'projects': projects_text
                }

                prompt = prompt_template.format(**context)

                achievements = self.ollama.generate_text(
                    prompt=prompt,
                    temperature=0.4,
                    max_tokens=300
                )

                if achievements and len(achievements.strip()) > 100:
                    self.logger.info("✓ Achievements paragraph generated successfully")
                    return achievements.strip()

            except Exception as e:
                self.logger.error(f"Error generating achievements: {e}")

        return None

    def _generate_closing(self, job, company_research: Dict[str, str]) -> Optional[str]:
        """Generate closing paragraph"""

        for attempt in range(1, self.max_retries + 1):
            try:
                self.logger.info(f"Generating closing (attempt {attempt}/{self.max_retries})")

                # Load prompt
                prompt_path = self.prompts_dir / 'cover_letter_sections' / 'closing.txt'
                with open(prompt_path, 'r', encoding='utf-8') as f:
                    prompt_template = f.read()

                context = {
                    'job_title': job.title,
                    'company_name': job.company,
                    'company_mission': company_research.get('mission', 'N/A'),
                    'job_description': job.description[:500]
                }

                prompt = prompt_template.format(**context)

                closing = self.ollama.generate_text(
                    prompt=prompt,
                    temperature=0.4,
                    max_tokens=150
                )

                if closing and len(closing.strip()) > 30:
                    self.logger.info("✓ Closing paragraph generated successfully")
                    return closing.strip()

            except Exception as e:
                self.logger.error(f"Error generating closing: {e}")

        return None

    def _fallback_opening(self, job) -> str:
        """Generic opening when research fails"""
        return f"I am writing to express my strong interest in the {job.title} position at {job.company}. As an MSc Data Science student with hands-on experience in machine learning and data analytics, I am excited about the opportunity to contribute to your team."

    def _fallback_achievements(self, matched_projects: List[Dict]) -> str:
        """Generic achievements using projects"""
        if not matched_projects:
            return "Through my academic projects and internship experience, I have developed strong skills in Python, machine learning, and data analysis. I have successfully delivered projects involving real-world data processing, predictive modeling, and automation systems."

        top_project = matched_projects[0]
        return f"In my recent project, {top_project['title']}, I demonstrated proficiency in {', '.join(top_project['technologies'][:3])}. Additionally, during my internship at RCCG Department of Public Health, I processed 500+ monthly transactions with 99.8% accuracy and reduced processing time by 25% through workflow optimization."

    def _fallback_closing(self, job) -> str:
        """Generic closing"""
        return f"I am eager to bring my technical skills and enthusiasm to {job.company}. I would welcome the opportunity to discuss how my background aligns with your team's goals. Thank you for your consideration."

    def _create_word_document(self, header: Dict[str, str], opening: str, achievements: str, closing: str, job) -> Document:
        """Create formatted Word document"""
        doc = Document()

        # Set margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)

        # Header - Name
        name_para = doc.add_paragraph(header['NAME'])
        if name_para.runs:
            name_para.runs[0].font.size = Pt(14)
            name_para.runs[0].bold = True

        # Header - Contact Line 1
        contact1 = f"{header['LOCATION']} | {header['PHONE']} | {header['EMAIL']}"
        contact1_para = doc.add_paragraph(contact1)
        if contact1_para.runs:
            contact1_para.runs[0].font.size = Pt(10)

        # Header - Contact Line 2
        contact2 = header['LINKEDIN']
        if header['GITHUB']:
            contact2 += f" | {header['GITHUB']}"
        contact2_para = doc.add_paragraph(contact2)
        if contact2_para.runs:
            contact2_para.runs[0].font.size = Pt(10)

        # Date
        doc.add_paragraph()
        date_para = doc.add_paragraph(header['DATE'])
        if date_para.runs:
            date_para.runs[0].font.size = Pt(11)

        # Recipient (if not Unknown)
        doc.add_paragraph()
        if job.company != "Unknown":
            recipient_para = doc.add_paragraph(f"Hiring Manager\n{job.company}")
            if recipient_para.runs:
                for run in recipient_para.runs:
                    run.font.size = Pt(11)

        # Salutation
        doc.add_paragraph()
        salutation = "Dear Hiring Manager," if job.company != "Unknown" else "Dear Hiring Team,"
        salutation_para = doc.add_paragraph(salutation)
        if salutation_para.runs:
            salutation_para.runs[0].font.size = Pt(11)

        # Opening paragraph
        doc.add_paragraph()
        opening_para = doc.add_paragraph(opening)
        if opening_para.runs:
            for run in opening_para.runs:
                run.font.size = Pt(11)

        # Achievements paragraph
        doc.add_paragraph()
        achievements_para = doc.add_paragraph(achievements)
        if achievements_para.runs:
            for run in achievements_para.runs:
                run.font.size = Pt(11)

        # Closing paragraph
        doc.add_paragraph()
        closing_para = doc.add_paragraph(closing)
        if closing_para.runs:
            for run in closing_para.runs:
                run.font.size = Pt(11)

        # Sign-off
        doc.add_paragraph()
        doc.add_paragraph("Best regards,")
        doc.add_paragraph(header['NAME'])

        return doc
